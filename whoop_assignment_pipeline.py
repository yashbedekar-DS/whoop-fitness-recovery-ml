from __future__ import annotations

import argparse
import math
import time
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from sklearn.compose import ColumnTransformer
from sklearn.dummy import DummyRegressor
from sklearn.ensemble import GradientBoostingRegressor, RandomForestRegressor
from sklearn.impute import SimpleImputer
from sklearn.inspection import permutation_importance
from sklearn.linear_model import ElasticNet, LinearRegression, Ridge
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from sklearn.model_selection import KFold, cross_validate, train_test_split
from sklearn.neighbors import KNeighborsRegressor
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.tree import DecisionTreeRegressor

try:
    from xgboost import XGBRegressor
except Exception:
    XGBRegressor = None


ROOT = Path(__file__).resolve().parent
DATA_PATH = ROOT / "data" / "whoop_fitness_dataset_100k.csv"
TABLE_DIR = ROOT / "outputs" / "tables"
IMAGE_DIR = ROOT / "outputs" / "figures" / "images"
NOTEBOOK_DIR = ROOT / "notebooks"
TARGET = "recovery_score"
DATASET_NAME = "WHOOP Fitness Dataset"
DATASET_LINK = "https://www.kaggle.com/datasets/likithagedipudi/whoop-fitness-dataset"
RANDOM_STATE = 42


def ensure_dirs() -> None:
    for path in [TABLE_DIR, IMAGE_DIR, NOTEBOOK_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def clean_outputs() -> None:
    ensure_dirs()
    for folder in [TABLE_DIR, IMAGE_DIR]:
        for file in folder.glob("*"):
            if file.is_file():
                file.unlink()


def encoder() -> OneHotEncoder:
    try:
        return OneHotEncoder(handle_unknown="ignore", sparse_output=False)
    except TypeError:
        return OneHotEncoder(handle_unknown="ignore", sparse=False)


def load_data() -> pd.DataFrame:
    if not DATA_PATH.exists():
        raise FileNotFoundError(f"Dataset not found: {DATA_PATH}")
    df = pd.read_csv(DATA_PATH)
    df["date"] = pd.to_datetime(df["date"], errors="coerce")
    df["date_month"] = df["date"].dt.month
    df["date_dayofyear"] = df["date"].dt.dayofyear
    return df


def modelling_data(sample_size: int = 30000) -> tuple[pd.DataFrame, pd.Series]:
    df = load_data()
    if sample_size and len(df) > sample_size:
        df = df.sample(sample_size, random_state=RANDOM_STATE).reset_index(drop=True)
    drop_cols = ["user_id", "date", TARGET]
    X = df.drop(columns=[c for c in drop_cols if c in df.columns])
    y = df[TARGET]
    return X, y


def columns_for(X: pd.DataFrame) -> tuple[list[str], list[str]]:
    categorical = X.select_dtypes(include=["object", "category"]).columns.tolist()
    numeric = [c for c in X.columns if c not in categorical]
    return numeric, categorical


def preprocessor_for(X: pd.DataFrame, scale_numeric: bool = True, categorical: bool = True) -> ColumnTransformer:
    numeric, cat_cols = columns_for(X)
    numeric_steps = [("imputer", SimpleImputer(strategy="median"))]
    if scale_numeric:
        numeric_steps.append(("scaler", StandardScaler()))
    transformers = [("num", Pipeline(numeric_steps), numeric)]
    if categorical and cat_cols:
        cat_pipe = Pipeline([
            ("imputer", SimpleImputer(strategy="most_frequent")),
            ("onehot", encoder()),
        ])
        transformers.append(("cat", cat_pipe, cat_cols))
    return ColumnTransformer(transformers=transformers, verbose_feature_names_out=False)


def make_pipeline(model, X: pd.DataFrame, scale_numeric: bool = True, categorical: bool = True) -> Pipeline:
    return Pipeline([
        ("preprocessor", preprocessor_for(X, scale_numeric=scale_numeric, categorical=categorical)),
        ("model", model),
    ])


def rmse(y_true, y_pred) -> float:
    return math.sqrt(mean_squared_error(y_true, y_pred))


def metrics_dict(y_true, y_pred) -> dict[str, float]:
    return {
        "MAE": mean_absolute_error(y_true, y_pred),
        "RMSE": rmse(y_true, y_pred),
        "R2": r2_score(y_true, y_pred),
    }


def rounded_report(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for col in out.select_dtypes(include=[np.number]).columns:
        out[col] = out[col].round(3)
    return out


def sampled(df: pd.DataFrame, size: int) -> pd.DataFrame:
    if len(df) <= size:
        return df
    return df.sample(size, random_state=RANDOM_STATE)


def error_cases(pred_df: pd.DataFrame, size: int = 20) -> pd.DataFrame:
    top = pred_df.nlargest(min(size, len(pred_df)), "Absolute_Error").copy()
    top["Error_Case"] = np.arange(1, len(top) + 1)
    return top


def permutation_sample(X: pd.DataFrame, y: pd.Series, size: int = 1000) -> tuple[pd.DataFrame, pd.Series]:
    X_sample = X.sample(min(size, len(X)), random_state=RANDOM_STATE)
    return X_sample, y.loc[X_sample.index]


def get_rankings_table() -> pd.DataFrame:
    rankings_path = TABLE_DIR / "RQ5_model_metric_rankings.csv"
    if rankings_path.exists():
        return pd.read_csv(rankings_path)
    return run_rq5()


def savefig(name: str) -> None:
    ensure_dirs()
    plt.tight_layout()
    plt.savefig(IMAGE_DIR / f"{name}.png", dpi=200, bbox_inches="tight")
    plt.close()


def split_data(sample_size: int = 30000):
    X, y = modelling_data(sample_size)
    return train_test_split(X, y, test_size=0.2, random_state=RANDOM_STATE), X, y


def baseline_models(X: pd.DataFrame) -> dict[str, Pipeline]:
    return {
        "Linear Regression": make_pipeline(LinearRegression(), X, scale_numeric=True),
        "Decision Tree": make_pipeline(
            DecisionTreeRegressor(max_depth=10, min_samples_leaf=30, random_state=RANDOM_STATE),
            X,
            scale_numeric=False,
        ),
        "k-NN": make_pipeline(KNeighborsRegressor(n_neighbors=15), X, scale_numeric=True),
    }


def candidate_models(X: pd.DataFrame) -> dict[str, Pipeline]:
    models = {
        "Linear Regression": make_pipeline(LinearRegression(), X, scale_numeric=True),
        "Ridge Regression": make_pipeline(Ridge(alpha=1.0), X, scale_numeric=True),
        "Elastic Net": make_pipeline(ElasticNet(alpha=0.01, l1_ratio=0.2, max_iter=5000), X, scale_numeric=True),
        "Decision Tree": make_pipeline(
            DecisionTreeRegressor(max_depth=12, min_samples_leaf=25, random_state=RANDOM_STATE),
            X,
            scale_numeric=False,
        ),
        "Random Forest": make_pipeline(
            RandomForestRegressor(
                n_estimators=80,
                max_depth=14,
                min_samples_leaf=20,
                random_state=RANDOM_STATE,
                n_jobs=-1,
            ),
            X,
            scale_numeric=False,
        ),
        "Gradient Boosting": make_pipeline(
            GradientBoostingRegressor(n_estimators=120, max_depth=3, random_state=RANDOM_STATE),
            X,
            scale_numeric=False,
        ),
    }
    if XGBRegressor is not None:
        models["XGBoost"] = make_pipeline(
            XGBRegressor(
                n_estimators=160,
                max_depth=4,
                learning_rate=0.06,
                subsample=0.9,
                colsample_bytree=0.9,
                random_state=RANDOM_STATE,
                objective="reg:squarederror",
                n_jobs=-1,
            ),
            X,
            scale_numeric=False,
        )
    models["Baseline Mean"] = make_pipeline(DummyRegressor(strategy="mean"), X, scale_numeric=False)
    return models


def evaluate_models(models: dict[str, Pipeline], X_train, X_test, y_train, y_test) -> tuple[pd.DataFrame, pd.DataFrame, dict]:
    rows = []
    pred_rows = []
    fitted = {}
    for name, pipe in models.items():
        start = time.perf_counter()
        pipe.fit(X_train, y_train)
        fit_seconds = time.perf_counter() - start
        start = time.perf_counter()
        preds = pipe.predict(X_test)
        predict_seconds = time.perf_counter() - start
        row = {"Model": name, **metrics_dict(y_test, preds), "Fit_Seconds": fit_seconds, "Predict_Seconds": predict_seconds}
        rows.append(row)
        fitted[name] = pipe
        sample_count = min(len(y_test), 1000)
        for actual, pred in zip(y_test.iloc[:sample_count], preds[:sample_count]):
            pred_rows.append({
                "Model": name,
                "Actual": actual,
                "Predicted": pred,
                "Residual": actual - pred,
                "Absolute_Error": abs(actual - pred),
            })
    return pd.DataFrame(rows), pd.DataFrame(pred_rows), fitted


def plot_error_metrics(table: pd.DataFrame, name: str, title: str) -> None:
    long = table.melt(id_vars="Model", value_vars=["MAE", "RMSE"], var_name="Metric", value_name="Value")
    plt.figure(figsize=(10, 5))
    sns.barplot(data=long, x="Model", y="Value", hue="Metric", palette=["#2a9d8f", "#e76f51"])
    plt.title(title)
    plt.xticks(rotation=25, ha="right")
    savefig(name)


def plot_actual_predicted(pred_df: pd.DataFrame, model: str, name: str, title: str) -> None:
    df = pred_df[pred_df["Model"] == model].copy()
    if len(df) > 800:
        df = df.sample(800, random_state=RANDOM_STATE)
    plt.figure(figsize=(6, 6))
    sns.scatterplot(data=df, x="Actual", y="Predicted", alpha=0.45, edgecolor=None, color="#2878b5")
    low = min(df["Actual"].min(), df["Predicted"].min())
    high = max(df["Actual"].max(), df["Predicted"].max())
    plt.plot([low, high], [low, high], color="#d62828", linestyle="--", linewidth=2)
    plt.title(title)
    savefig(name)


def plot_residuals(pred_df: pd.DataFrame, name: str, title: str) -> None:
    df = pred_df.copy()
    if len(df) > 2500:
        df = df.sample(2500, random_state=RANDOM_STATE)
    plt.figure(figsize=(10, 5))
    sns.boxplot(data=df, x="Model", y="Residual", color="#8ecae6")
    plt.axhline(0, color="#d62828", linestyle="--")
    plt.title(title)
    plt.xticks(rotation=25, ha="right")
    savefig(name)


def run_rq1() -> pd.DataFrame:
    ensure_dirs()
    (X_train, X_test, y_train, y_test), X, _ = split_data(30000)
    table, preds, fitted = evaluate_models(baseline_models(X), X_train, X_test, y_train, y_test)
    table = table[["Model", "MAE", "RMSE", "R2"]]
    table["Expected_R2_Range"] = "0.45-0.75"
    table["Matches_Expected_Range"] = table["R2"].between(0.45, 0.75).map({True: "Yes", False: "No"})
    table["Interpretation"] = np.where(
        table["R2"] > 0.75,
        "Above expected range; strong recovery signal in available fitness features.",
        np.where(table["R2"] >= 0.45, "Matches expected moderate baseline performance.", "Below expected range."),
    )
    table.to_csv(TABLE_DIR / "RQ1_table.csv", index=False)
    rounded_report(table).to_csv(TABLE_DIR / "RQ1_table_report_rounded.csv", index=False)
    preds.to_csv(TABLE_DIR / "RQ1_predictions.csv", index=False)
    plot_error_metrics(table, "RQ1_error_metrics", "RQ1: Baseline Error Metrics")
    best = table.sort_values("R2", ascending=False).iloc[0]["Model"]
    plot_actual_predicted(preds, best, "RQ1_actual_vs_predicted", f"RQ1: Actual vs Predicted ({best})")
    plot_residuals(preds, "RQ1_residuals_by_model", "RQ1: Residuals by Baseline Model")
    plt.figure(figsize=(8, 4))
    sns.barplot(data=table, x="Model", y="R2", color="#457b9d")
    plt.axhspan(0.45, 0.75, color="#a8dadc", alpha=0.3, label="Expected range")
    plt.ylim(0, max(1.0, table["R2"].max() + 0.05))
    plt.title("RQ1: R2 Values Against Expected Baseline Range")
    plt.xticks(rotation=20, ha="right")
    plt.legend()
    savefig("RQ1_r2_expected_range")
    df = load_data()
    corr = df.select_dtypes(include=[np.number]).corr(numeric_only=True)[TARGET].drop(TARGET).abs().sort_values(ascending=False).head(12)
    corr_df = corr.reset_index()
    corr_df.columns = ["Feature", "Absolute_Correlation_With_Recovery"]
    corr_df.to_csv(TABLE_DIR / "RQ1_top_numeric_correlations.csv", index=False)
    plt.figure(figsize=(8, 5))
    sns.barplot(data=corr_df, y="Feature", x="Absolute_Correlation_With_Recovery", color="#2a9d8f")
    plt.title("RQ1: Top Numeric Correlations with Recovery Score")
    savefig("RQ1_top_numeric_correlations")
    return table


def run_rq2() -> pd.DataFrame:
    ensure_dirs()
    (X_train, X_test, y_train, y_test), X, _ = split_data(30000)
    models = candidate_models(X)
    table, preds, fitted = evaluate_models(models, X_train, X_test, y_train, y_test)
    table = table[table["Model"] != "Baseline Mean"].copy()
    table["Model_Type"] = table["Model"].map(lambda m: "Ensemble" if m in ["Random Forest", "Gradient Boosting", "XGBoost"] else "Simple/regularized baseline")
    table["Rank_by_R2"] = table["R2"].rank(ascending=False, method="min").astype(int)
    table["Rank_by_MAE"] = table["MAE"].rank(ascending=True, method="min").astype(int)
    best = table.sort_values("R2", ascending=False).iloc[0]["Model"]
    table["Best_Model"] = np.where(table["Model"] == best, "Yes", "No")
    table["XGBoost_Hypothesis_Result"] = f"{'Supported' if best == 'XGBoost' else 'Not supported'}; {best} achieved the best R2 performance."
    table = table[["Model", "Model_Type", "MAE", "RMSE", "R2", "Rank_by_R2", "Rank_by_MAE", "Best_Model", "XGBoost_Hypothesis_Result"]]
    table.to_csv(TABLE_DIR / "RQ2_table.csv", index=False)
    rounded_report(table).to_csv(TABLE_DIR / "RQ2_table_report_rounded.csv", index=False)
    preds.to_csv(TABLE_DIR / "RQ2_predictions.csv", index=False)
    plot_error_metrics(table, "RQ2_error_metrics", "RQ2: Candidate Model Error Metrics")
    plt.figure(figsize=(9, 5))
    sns.barplot(data=table.sort_values("R2", ascending=False), x="R2", y="Model", hue="Model_Type", dodge=False)
    plt.title("RQ2: R2 Comparison of Candidate Models")
    savefig("RQ2_R2_comparison")
    plot_actual_predicted(preds, best, "RQ2_actual_vs_predicted", f"RQ2: Actual vs Predicted ({best})")
    simple_best = table[table["Model_Type"] != "Ensemble"].sort_values("R2", ascending=False).iloc[0]
    ensemble_best = table[table["Model_Type"] == "Ensemble"].sort_values("R2", ascending=False).iloc[0]
    gain = pd.DataFrame([{
        "Best_Simple_Model": simple_best["Model"],
        "Best_Simple_R2": simple_best["R2"],
        "Best_Ensemble_Model": ensemble_best["Model"],
        "Best_Ensemble_R2": ensemble_best["R2"],
        "R2_Gain": ensemble_best["R2"] - simple_best["R2"],
        "MAE_Reduction": simple_best["MAE"] - ensemble_best["MAE"],
    }])
    gain.to_csv(TABLE_DIR / "RQ2_ensemble_gain_summary.csv", index=False)
    plt.figure(figsize=(6, 4))
    sns.barplot(data=gain.melt(value_vars=["Best_Simple_R2", "Best_Ensemble_R2"], var_name="Group", value_name="R2"), x="Group", y="R2", palette=["#e76f51", "#2a9d8f"], hue="Group", legend=False)
    plt.title("RQ2: Ensemble Gain Over Best Simple Model")
    savefig("RQ2_ensemble_gain")
    best_pipe = fitted[best]
    perm_sample, y_perm = permutation_sample(X_test, y_test)
    perm = permutation_importance(best_pipe, perm_sample, y_perm, n_repeats=5, random_state=RANDOM_STATE, scoring="r2")
    perm_df = pd.DataFrame({"Feature": X.columns, "Mean_R2_Decrease": perm.importances_mean, "Std_R2_Decrease": perm.importances_std}).sort_values("Mean_R2_Decrease", ascending=False).head(15)
    perm_df.to_csv(TABLE_DIR / "RQ2_best_model_permutation_importance.csv", index=False)
    plt.figure(figsize=(8, 5))
    sns.barplot(data=perm_df, y="Feature", x="Mean_R2_Decrease", color="#6a994e")
    plt.title(f"RQ2: Permutation Importance ({best})")
    savefig("RQ2_best_model_permutation_importance")
    return table


def run_rq3() -> pd.DataFrame:
    ensure_dirs()
    X, y = modelling_data(30000)
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=RANDOM_STATE)
    numeric, _ = columns_for(X)
    strategies = {
        "Numeric only + imputation": Pipeline([
            ("preprocessor", ColumnTransformer([("num", SimpleImputer(strategy="median"), numeric)])),
            ("model", Ridge(alpha=1.0)),
        ]),
        "Numeric + imputation + scaling": Pipeline([
            ("preprocessor", ColumnTransformer([("num", Pipeline([("imputer", SimpleImputer(strategy="median")), ("scaler", StandardScaler())]), numeric)])),
            ("model", Ridge(alpha=1.0)),
        ]),
        "Full pipeline: scaling + encoding": make_pipeline(Ridge(alpha=1.0), X, scale_numeric=True, categorical=True),
        "Full pipeline + Random Forest": make_pipeline(
            RandomForestRegressor(n_estimators=80, max_depth=14, min_samples_leaf=20, random_state=RANDOM_STATE, n_jobs=-1),
            X,
            scale_numeric=False,
            categorical=True,
        ),
    }
    rows = []
    pred_rows = []
    for name, pipe in strategies.items():
        pipe.fit(X_train, y_train)
        preds = pipe.predict(X_test)
        rows.append({"Preprocessing_Strategy": name, **metrics_dict(y_test, preds)})
        if name == "Full pipeline + Random Forest":
            for actual, pred in zip(y_test.iloc[:1000], preds[:1000]):
                pred_rows.append({"Actual": actual, "Predicted": pred, "Residual": actual - pred, "Absolute_Error": abs(actual - pred)})
    table = pd.DataFrame(rows)
    table.to_csv(TABLE_DIR / "RQ3_table.csv", index=False)
    pd.DataFrame(pred_rows).to_csv(TABLE_DIR / "RQ3_predictions.csv", index=False)
    plt.figure(figsize=(9, 5))
    sns.barplot(data=table.sort_values("R2", ascending=False), x="R2", y="Preprocessing_Strategy", color="#457b9d")
    plt.title("RQ3: Effect of Preprocessing on R2")
    savefig("RQ3_R2_comparison")
    pred_df = pd.DataFrame(pred_rows)
    plt.figure(figsize=(6, 6))
    sns.scatterplot(data=sampled(pred_df, 800), x="Actual", y="Predicted", alpha=0.45, edgecolor=None)
    low = min(pred_df["Actual"].min(), pred_df["Predicted"].min())
    high = max(pred_df["Actual"].max(), pred_df["Predicted"].max())
    plt.plot([low, high], [low, high], color="#d62828", linestyle="--")
    plt.title("RQ3: Actual vs Predicted (Full Pipeline + Random Forest)")
    savefig("RQ3_actual_vs_predicted")
    return table


def original_feature_name(encoded_name: str, original_cols: list[str]) -> str:
    if encoded_name in original_cols:
        return encoded_name
    for col in original_cols:
        if encoded_name.startswith(f"{col}_"):
            return col
    return encoded_name


def run_rq4() -> pd.DataFrame:
    ensure_dirs()
    (X_train, X_test, y_train, y_test), X, _ = split_data(30000)
    pipe = make_pipeline(
        RandomForestRegressor(n_estimators=100, max_depth=14, min_samples_leaf=20, random_state=RANDOM_STATE, n_jobs=-1),
        X,
        scale_numeric=False,
    )
    pipe.fit(X_train, y_train)
    preds = pipe.predict(X_test)
    performance = pd.DataFrame([
        {"Metric": "MAE", "Value": mean_absolute_error(y_test, preds)},
        {"Metric": "RMSE", "Value": rmse(y_test, preds)},
        {"Metric": "R2", "Value": r2_score(y_test, preds)},
        {"Metric": "Train Rows", "Value": len(X_train)},
        {"Metric": "Test Rows", "Value": len(X_test)},
    ])
    pre = pipe.named_steps["preprocessor"]
    rf = pipe.named_steps["model"]
    encoded_names = pre.get_feature_names_out()
    feature_df = pd.DataFrame({"Feature": encoded_names, "Importance": rf.feature_importances_}).sort_values("Importance", ascending=False)
    feature_df.insert(0, "Rank", np.arange(1, len(feature_df) + 1))
    feature_df["Importance_Percent"] = feature_df["Importance"] * 100
    feature_df["Original_Feature"] = feature_df["Feature"].map(lambda x: original_feature_name(x, X.columns.tolist()))
    agg = feature_df.groupby("Original_Feature", as_index=False)["Importance"].sum().sort_values("Importance", ascending=False)
    agg.insert(0, "Rank", np.arange(1, len(agg) + 1))
    agg["Importance_Percent"] = agg["Importance"] * 100
    performance.to_csv(TABLE_DIR / "RQ4_model_performance.csv", index=False)
    feature_df.to_csv(TABLE_DIR / "RQ4_feature_importance.csv", index=False)
    agg.to_csv(TABLE_DIR / "RQ4_aggregated_feature_importance.csv", index=False)
    expected = agg[agg["Original_Feature"].isin(["hrv", "resting_heart_rate", "sleep_hours", "day_strain", "sleep_efficiency", "calories_burned"])]
    expected.to_csv(TABLE_DIR / "RQ4_expected_feature_check.csv", index=False)
    pred_table = X_test.copy()
    pred_table["Actual_Recovery_Score"] = y_test
    pred_table["Predicted_Recovery_Score"] = preds
    pred_table["Residual"] = y_test - preds
    pred_table["Absolute_Error"] = np.abs(pred_table["Residual"])
    pred_table.head(2000).to_csv(TABLE_DIR / "RQ4_predictions_table.csv", index=False)
    plot_actual_predicted(pd.DataFrame({"Model": "Random Forest", "Actual": y_test, "Predicted": preds}), "Random Forest", "RQ4_actual_vs_predicted", "RQ4: Actual vs Predicted (Random Forest)")
    residual_df = pd.DataFrame({"Actual": y_test, "Predicted": preds, "Residual": y_test - preds})
    plt.figure(figsize=(7, 5))
    sns.scatterplot(data=sampled(residual_df, 1000), x="Predicted", y="Residual", alpha=0.45, edgecolor=None)
    plt.axhline(0, color="#d62828", linestyle="--")
    plt.title("RQ4: Residuals for Random Forest")
    savefig("RQ4_residuals")
    plt.figure(figsize=(9, 6))
    sns.barplot(data=agg.head(15), y="Original_Feature", x="Importance_Percent", color="#2a9d8f")
    plt.title("RQ4: Top Aggregated Feature Importances")
    savefig("RQ4_feature_importance_bar")
    numeric_corr = load_data().select_dtypes(include=[np.number]).corr(numeric_only=True)
    top = numeric_corr[TARGET].abs().sort_values(ascending=False).head(12).index
    plt.figure(figsize=(9, 7))
    sns.heatmap(numeric_corr.loc[top, top], cmap="vlag", center=0)
    plt.title("RQ4: Correlation Heatmap of Top Numeric Recovery Features")
    savefig("RQ4_correlation_heatmap")
    perm_sample, y_perm = permutation_sample(X_test, y_test)
    perm = permutation_importance(pipe, perm_sample, y_perm, n_repeats=5, random_state=RANDOM_STATE, scoring="r2")
    perm_df = pd.DataFrame({"Feature": X.columns, "Mean_R2_Decrease": perm.importances_mean, "Std_R2_Decrease": perm.importances_std}).sort_values("Mean_R2_Decrease", ascending=False)
    perm_df.to_csv(TABLE_DIR / "RQ4_permutation_importance.csv", index=False)
    plt.figure(figsize=(8, 5))
    sns.barplot(data=perm_df.head(15), y="Feature", x="Mean_R2_Decrease", color="#6a994e")
    plt.title("RQ4: Permutation Importance")
    savefig("RQ4_permutation_importance")
    return agg


def run_rq5() -> pd.DataFrame:
    ensure_dirs()
    X, y = modelling_data(20000)
    models = candidate_models(X)
    cv = KFold(n_splits=5, shuffle=True, random_state=RANDOM_STATE)
    scoring = {"mae": "neg_mean_absolute_error", "rmse": "neg_root_mean_squared_error", "r2": "r2"}
    rows = []
    for name, pipe in models.items():
        scores = cross_validate(pipe, X, y, cv=cv, scoring=scoring, n_jobs=None)
        rows.append({
            "Model": name,
            "MAE": -scores["test_mae"].mean(),
            "MAE_Std": scores["test_mae"].std(),
            "RMSE": -scores["test_rmse"].mean(),
            "RMSE_Std": scores["test_rmse"].std(),
            "R2": scores["test_r2"].mean(),
            "R2_Std": scores["test_r2"].std(),
        })
    table = pd.DataFrame(rows)
    table["MAE_Rank"] = table["MAE"].rank(ascending=True, method="min").astype(int)
    table["RMSE_Rank"] = table["RMSE"].rank(ascending=True, method="min").astype(int)
    table["R2_Rank"] = table["R2"].rank(ascending=False, method="min").astype(int)
    table["Average_Rank"] = table[["MAE_Rank", "RMSE_Rank", "R2_Rank"]].mean(axis=1)
    table["Rank_Range"] = table[["MAE_Rank", "RMSE_Rank", "R2_Rank"]].max(axis=1) - table[["MAE_Rank", "RMSE_Rank", "R2_Rank"]].min(axis=1)
    table = table.sort_values("Average_Rank")
    table.to_csv(TABLE_DIR / "RQ5_model_metric_rankings.csv", index=False)
    rankings_long = table.melt(id_vars="Model", value_vars=["MAE_Rank", "RMSE_Rank", "R2_Rank"], var_name="Metric", value_name="Rank")
    rankings_long.to_csv(TABLE_DIR / "RQ5_rankings_long.csv", index=False)
    stability = table[["Model", "Average_Rank", "Rank_Range"]].copy()
    stability.to_csv(TABLE_DIR / "RQ5_ranking_stability.csv", index=False)
    plt.figure(figsize=(10, 5))
    sns.barplot(data=table.melt(id_vars="Model", value_vars=["MAE", "RMSE"], var_name="Metric", value_name="Value"), x="Model", y="Value", hue="Metric")
    plt.xticks(rotation=25, ha="right")
    plt.title("RQ5: Metric Comparison by Model")
    savefig("RQ5_metric_comparison_bars")
    heat = table.set_index("Model")[["MAE_Rank", "RMSE_Rank", "R2_Rank"]]
    plt.figure(figsize=(8, 5))
    sns.heatmap(heat, annot=True, cmap="YlGnBu_r", cbar_kws={"label": "Rank"})
    plt.title("RQ5: Rank Heatmap")
    savefig("RQ5_rank_heatmap")
    plt.figure(figsize=(9, 5))
    for model, grp in rankings_long.groupby("Model"):
        plt.plot(grp["Metric"], grp["Rank"], marker="o", label=model)
    plt.gca().invert_yaxis()
    plt.title("RQ5: Ranking Bump Chart")
    plt.ylabel("Rank")
    plt.legend(bbox_to_anchor=(1.02, 1), loc="upper left")
    savefig("RQ5_ranking_bump_chart")
    best = table.iloc[0]["Model"]
    X_full, y_full = modelling_data(30000)
    X_train, X_test, y_train, y_test = train_test_split(X_full, y_full, test_size=0.2, random_state=RANDOM_STATE)
    best_pipe = candidate_models(X_full)[best]
    best_pipe.fit(X_train, y_train)
    preds = best_pipe.predict(X_test)
    holdout = pd.DataFrame([{"Metric": "Best Model", "Value": best}, {"Metric": "Holdout MAE", "Value": mean_absolute_error(y_test, preds)}, {"Metric": "Holdout RMSE", "Value": rmse(y_test, preds)}, {"Metric": "Holdout R2", "Value": r2_score(y_test, preds)}])
    holdout.to_csv(TABLE_DIR / "RQ5_best_model_holdout_summary.csv", index=False)
    pred_df = pd.DataFrame({"Actual": y_test, "Predicted": preds, "Residual": y_test - preds, "Absolute_Error": np.abs(y_test - preds)})
    pred_df.head(2000).to_csv(TABLE_DIR / "RQ5_best_model_predictions.csv", index=False)
    plot_actual_predicted(pd.DataFrame({"Model": best, "Actual": y_test, "Predicted": preds}), best, "RQ5_best_model_actual_vs_predicted", f"RQ5: Actual vs Predicted ({best})")
    top_errors = error_cases(pred_df)
    plt.figure(figsize=(8, 5))
    sns.barplot(data=top_errors, x="Absolute_Error", y="Error_Case", color="#e76f51", orient="h")
    plt.ylabel("Largest Error Cases")
    plt.title("RQ5: Largest Prediction Errors")
    savefig("RQ5_largest_prediction_errors")
    return table


def cv_summary(pipe: Pipeline, X: pd.DataFrame, y: pd.Series, label: str) -> dict:
    cv = KFold(n_splits=3, shuffle=True, random_state=RANDOM_STATE)
    scores = cross_validate(pipe, X, y, cv=cv, scoring={"r2": "r2", "rmse": "neg_root_mean_squared_error", "mae": "neg_mean_absolute_error"})
    return {
        "Condition": label,
        "Mean_R2": scores["test_r2"].mean(),
        "Min_R2": scores["test_r2"].min(),
        "Max_R2": scores["test_r2"].max(),
        "Mean_RMSE": -scores["test_rmse"].mean(),
        "Mean_MAE": -scores["test_mae"].mean(),
    }


def run_rq6() -> pd.DataFrame:
    ensure_dirs()
    rankings = get_rankings_table()
    best = rankings.sort_values("R2", ascending=False).iloc[0]["Model"]
    X, y = modelling_data(15000)
    model = candidate_models(X)[best]
    results = [cv_summary(model, X, y, "Baseline CV")]
    numeric, categorical = columns_for(X)
    noisy = X.copy()
    rng = np.random.default_rng(RANDOM_STATE)
    for col in numeric:
        std = noisy[col].std()
        if std and not np.isnan(std):
            noisy[col] = noisy[col] + rng.normal(0, std * 0.05, len(noisy))
    results.append(cv_summary(candidate_models(noisy)[best], noisy, y, "Added numeric noise"))
    missing = X.copy()
    mask = rng.random(missing.shape) < 0.10
    missing = missing.mask(mask)
    results.append(cv_summary(candidate_models(missing)[best], missing, y, "Simulated missingness"))
    robustness = pd.DataFrame(results)
    robustness.to_csv(TABLE_DIR / "RQ6_robustness_summary.csv", index=False)
    X_full, y_full = modelling_data(30000)
    X_train, X_test, y_train, y_test = train_test_split(X_full, y_full, test_size=0.2, random_state=RANDOM_STATE)
    holdout_pipe = candidate_models(X_full)[best]
    holdout_pipe.fit(X_train, y_train)
    preds = holdout_pipe.predict(X_test)
    holdout = pd.DataFrame([{"Metric": "Model", "Value": best}, {"Metric": "Holdout MAE", "Value": mean_absolute_error(y_test, preds)}, {"Metric": "Holdout RMSE", "Value": rmse(y_test, preds)}, {"Metric": "Holdout R2", "Value": r2_score(y_test, preds)}])
    holdout.to_csv(TABLE_DIR / "RQ6_holdout_metrics.csv", index=False)
    pred_df = pd.DataFrame({"Actual": y_test, "Predicted": preds, "Residual": y_test - preds, "Absolute_Error": np.abs(y_test - preds)})
    pred_df.head(2000).to_csv(TABLE_DIR / "RQ6_holdout_predictions.csv", index=False)
    robust_long = robustness.melt(id_vars="Condition", value_vars=["Mean_R2", "Mean_RMSE", "Mean_MAE"], var_name="Metric", value_name="Value")
    robust_long.to_csv(TABLE_DIR / "RQ6_robustness_results.csv", index=False)
    plot_actual_predicted(pd.DataFrame({"Model": best, "Actual": y_test, "Predicted": preds}), best, "RQ6_holdout_actual_vs_predicted", f"RQ6: Holdout Actual vs Predicted ({best})")
    plt.figure(figsize=(8, 5))
    sns.barplot(data=robustness, x="Condition", y="Mean_R2", color="#457b9d")
    plt.title("RQ6: R2 Robustness")
    plt.xticks(rotation=20, ha="right")
    savefig("RQ6_r2_robustness_bar")
    plt.figure(figsize=(8, 5))
    sns.barplot(data=robustness, x="Condition", y="Mean_RMSE", color="#e76f51")
    plt.title("RQ6: RMSE Robustness")
    plt.xticks(rotation=20, ha="right")
    savefig("RQ6_rmse_robustness_bar")
    base = robustness.loc[robustness["Condition"] == "Baseline CV", "Mean_R2"].iloc[0]
    degradation = robustness.copy()
    degradation["R2_Degradation"] = base - degradation["Mean_R2"]
    plt.figure(figsize=(8, 5))
    sns.barplot(data=degradation, x="Condition", y="R2_Degradation", color="#f4a261")
    plt.title("RQ6: R2 Degradation From Baseline")
    plt.xticks(rotation=20, ha="right")
    savefig("RQ6_r2_degradation")
    top_errors = error_cases(pred_df)
    plt.figure(figsize=(8, 5))
    sns.barplot(data=top_errors, x="Absolute_Error", y="Error_Case", color="#e76f51", orient="h")
    plt.ylabel("Largest Error Cases")
    plt.title("RQ6: Largest Prediction Errors")
    savefig("RQ6_largest_prediction_errors")
    return robustness


def score_1_to_5(series: pd.Series, higher_better: bool = True) -> pd.Series:
    s = series.astype(float)
    if s.max() == s.min():
        return pd.Series(3.0, index=s.index)
    scaled = (s - s.min()) / (s.max() - s.min())
    if not higher_better:
        scaled = 1 - scaled
    return 1 + 4 * scaled


def run_rq7() -> pd.DataFrame:
    ensure_dirs()
    rankings = get_rankings_table()
    X, y = modelling_data(30000)
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=RANDOM_STATE)
    rows = []
    for name in rankings["Model"]:
        pipe = candidate_models(X)[name]
        start = time.perf_counter()
        pipe.fit(X_train, y_train)
        fit_seconds = time.perf_counter() - start
        start = time.perf_counter()
        preds = pipe.predict(X_test)
        predict_seconds = time.perf_counter() - start
        cv_row = rankings[rankings["Model"] == name].iloc[0]
        rows.append({
            "Model": name,
            "CV_MAE": cv_row["MAE"],
            "CV_RMSE": cv_row["RMSE"],
            "CV_R2": cv_row["R2"],
            "CV_R2_Std": cv_row["R2_Std"],
            "Holdout_MAE": mean_absolute_error(y_test, preds),
            "Holdout_RMSE": rmse(y_test, preds),
            "Holdout_R2": r2_score(y_test, preds),
            "Fit_Seconds": fit_seconds,
            "Predict_Seconds": predict_seconds,
        })
    table = pd.DataFrame(rows)
    interpretability = {
        "Linear Regression": 5,
        "Ridge Regression": 5,
        "Elastic Net": 5,
        "Decision Tree": 4,
        "Random Forest": 3.5,
        "Gradient Boosting": 3,
        "XGBoost": 2.5,
        "Baseline Mean": 1,
    }
    deployment = {
        "Linear Regression": 5,
        "Ridge Regression": 5,
        "Elastic Net": 5,
        "Decision Tree": 4,
        "Random Forest": 4,
        "Gradient Boosting": 3,
        "XGBoost": 3,
        "Baseline Mean": 2,
    }
    table["Interpretability_Score"] = table["Model"].map(interpretability).fillna(3)
    table["Deployment_Suitability_Score"] = table["Model"].map(deployment).fillna(3)
    table["Accuracy_Score"] = score_1_to_5(table["Holdout_R2"], True)
    table["Reliability_Score"] = score_1_to_5(table["CV_R2_Std"], False)
    table["Computational_Efficiency_Score"] = score_1_to_5(table["Fit_Seconds"] + table["Predict_Seconds"], False)
    table["Weighted_Decision_Score"] = (
        0.35 * table["Accuracy_Score"]
        + 0.20 * table["Reliability_Score"]
        + 0.20 * table["Interpretability_Score"]
        + 0.15 * table["Deployment_Suitability_Score"]
        + 0.10 * table["Computational_Efficiency_Score"]
    )
    table["Final_Rank"] = table["Weighted_Decision_Score"].rank(ascending=False, method="min").astype(int)
    table = table.sort_values("Final_Rank")
    table.to_csv(TABLE_DIR / "RQ7_decision_matrix.csv", index=False)
    table[["Model", "CV_R2", "Holdout_R2", "CV_R2_Std", "Weighted_Decision_Score", "Final_Rank"]].to_csv(TABLE_DIR / "RQ7_model_performance_reliability.csv", index=False)
    recommended = table.iloc[0]["Model"]
    highest_holdout = table.sort_values("Holdout_R2", ascending=False).iloc[0]["Model"]
    final = pd.DataFrame([
        {"Item": "Recommended Model", "Value": recommended},
        {"Item": "Highest Holdout Accuracy Model", "Value": highest_holdout},
        {"Item": "Recommendation", "Value": "Use the recommended model for planning because it balances accuracy, reliability, interpretability, runtime, and deployment suitability."},
    ])
    final.to_csv(TABLE_DIR / "RQ7_final_recommendation.csv", index=False)
    weights = pd.DataFrame([
        {"Criterion": "Accuracy", "Weight": 0.35},
        {"Criterion": "Reliability", "Weight": 0.20},
        {"Criterion": "Interpretability", "Weight": 0.20},
        {"Criterion": "Deployment suitability", "Weight": 0.15},
        {"Criterion": "Computational efficiency", "Weight": 0.10},
    ])
    weights.to_csv(TABLE_DIR / "RQ7_decision_criteria_weights.csv", index=False)
    rec_pipe = candidate_models(X)[recommended]
    rec_pipe.fit(X_train, y_train)
    preds = rec_pipe.predict(X_test)
    pred_df = pd.DataFrame({"Actual": y_test, "Predicted": preds, "Residual": y_test - preds, "Absolute_Error": np.abs(y_test - preds)})
    pred_df.head(2000).to_csv(TABLE_DIR / "RQ7_recommended_model_predictions.csv", index=False)
    perm_sample, y_perm = permutation_sample(X_test, y_test)
    perm = permutation_importance(rec_pipe, perm_sample, y_perm, n_repeats=5, random_state=RANDOM_STATE, scoring="r2")
    perm_df = pd.DataFrame({"Feature": X.columns, "Mean_R2_Decrease": perm.importances_mean, "Std_R2_Decrease": perm.importances_std}).sort_values("Mean_R2_Decrease", ascending=False)
    perm_df.to_csv(TABLE_DIR / "RQ7_recommended_model_permutation_importance.csv", index=False)
    heat_cols = ["Accuracy_Score", "Reliability_Score", "Interpretability_Score", "Deployment_Suitability_Score", "Computational_Efficiency_Score"]
    plt.figure(figsize=(10, 6))
    sns.heatmap(table.set_index("Model")[heat_cols], annot=True, fmt=".2f", cmap="YlGnBu")
    plt.title("RQ7: Final Decision Matrix")
    savefig("RQ7_decision_matrix_heatmap")
    plt.figure(figsize=(8, 5))
    sns.barplot(data=table, x="Weighted_Decision_Score", y="Model", color="#2a9d8f")
    plt.title("RQ7: Weighted Decision Score")
    savefig("RQ7_weighted_decision_score")
    plt.figure(figsize=(8, 5))
    sns.scatterplot(data=table, x="Holdout_R2", y="Interpretability_Score", size="Reliability_Score", hue="Model", sizes=(80, 400))
    plt.title("RQ7: Performance, Interpretability, and Reliability")
    plt.legend(bbox_to_anchor=(1.02, 1), loc="upper left")
    savefig("RQ7_performance_comparison")
    plot_actual_predicted(pd.DataFrame({"Model": recommended, "Actual": y_test, "Predicted": preds}), recommended, "RQ7_recommended_actual_vs_predicted", f"RQ7: Actual vs Predicted ({recommended})")
    plt.figure(figsize=(8, 5))
    sns.barplot(data=perm_df.head(15), y="Feature", x="Mean_R2_Decrease", color="#6a994e")
    plt.title(f"RQ7: Recommended Model Permutation Importance ({recommended})")
    savefig("RQ7_recommended_permutation_importance")
    top_errors = error_cases(pred_df)
    plt.figure(figsize=(8, 5))
    sns.barplot(data=top_errors, x="Absolute_Error", y="Error_Case", color="#e76f51", orient="h")
    plt.ylabel("Largest Error Cases")
    plt.title("RQ7: Largest Prediction Errors")
    savefig("RQ7_largest_prediction_errors")
    return table


def write_markdown_files() -> None:
    ensure_dirs()
    df = load_data()
    final = pd.read_csv(TABLE_DIR / "RQ7_final_recommendation.csv")
    recommended = final.loc[final["Item"] == "Recommended Model", "Value"].iloc[0]
    readme = f"""# Submission 1: Proposal, Code, and GitHub Repository

## Project Title

Supervised Learning for WHOOP Fitness Recovery Score Prediction

## Dataset Link

Original public dataset: [{DATASET_NAME}]({DATASET_LINK})

Local dataset file:

```text
data/whoop_fitness_dataset_100k.csv
```

## Project Overview

This project uses supervised machine learning to predict `recovery_score` from WHOOP-style fitness, sleep, heart-rate, strain, workout, and user profile indicators.

## Dataset Details

| Item | Value |
|---|---|
| Dataset name | {DATASET_NAME} |
| Source | Kaggle |
| Rows | {df.shape[0]:,} |
| Columns | {df.shape[1]} |
| Target variable | `recovery_score` |
| Task type | Regression |
| Application domain | Fitness, recovery, and wearable health analytics |

## Research Questions

1. How effectively can baseline supervised learning models predict WHOOP recovery score?
2. Which supervised learning model achieves the best predictive performance for recovery score?
3. How do preprocessing strategies affect model performance?
4. Which input features contribute most to recovery score prediction?
5. How does model ranking change across MAE, RMSE, and R2?
6. How robust is the best model under cross-validation, noise, and missing-data conditions?
7. Which model provides the best practical balance of accuracy, interpretability, reliability, runtime, and deployment suitability?

## How To Run The Code

Install the required libraries:

```bash
python -m pip install -r requirements.txt
```

Run all outputs:

```bash
python whoop_assignment_pipeline.py --run-all
```

Windows launcher alternative:

```bash
py -3 -m pip install -r requirements.txt
py -3 whoop_assignment_pipeline.py --run-all
```

Or open and run the notebooks in `notebooks/` from `RQ1_notebook.ipynb` to `RQ7_notebook.ipynb`.

## Generated Outputs

Tables are saved in `outputs/tables/`. Figures are saved in `outputs/figures/images/`.

## Key Final Result

The final recommended model is `{recommended}` because it provides the best weighted balance across accuracy, reliability, interpretability, runtime, and deployment suitability.

## Submission Checklist

- Proposal file: `Submission_1_Proposal.docx`
- Markdown proposal: `Submission_1_Proposal.md`
- Complete code: `whoop_assignment_pipeline.py` and `notebooks/RQ1_notebook.ipynb` to `notebooks/RQ7_notebook.ipynb`
- README file: `README.md`
- Dataset link: {DATASET_LINK}
- Generated outputs: `outputs/`
- Required libraries: `requirements.txt`
"""
    (ROOT / "README.md").write_text(readme, encoding="utf-8")
    proposal = f"""# Submission 1 Proposal

## Project Title

Supervised Learning for WHOOP Fitness Recovery Score Prediction

## Dataset Name, Source Link, Rows, And Columns

| Item | Details |
|---|---|
| Dataset name | {DATASET_NAME} |
| Source link | {DATASET_LINK} |
| Local file | `data/whoop_fitness_dataset_100k.csv` |
| Number of rows | {df.shape[0]:,} |
| Number of columns | {df.shape[1]} |

## Target Variable

`recovery_score`

## Type Of Task

Regression. The target is a continuous recovery/readiness score.

## Problem Statement

Wearable fitness platforms collect daily sleep, strain, workout, heart-rate, and physiological measurements. This project builds supervised regression models to predict WHOOP recovery score and identify which wearable indicators are most useful for recovery planning and fitness decision-making.

## Research Questions

1. How effectively can baseline supervised learning models predict WHOOP recovery score?
2. Which supervised learning model achieves the best predictive performance for recovery score?
3. How do preprocessing strategies affect supervised learning performance?
4. Which features contribute most to recovery score prediction?
5. How does model ranking change across MAE, RMSE, and R2?
6. How robust is the best supervised learning model under cross-validation, numeric noise, and simulated missingness?
7. Which model is most useful, interpretable, reliable, and practical for recovery-score decision support?

## Proposed Methodology

1. Load the WHOOP Fitness dataset.
2. Use `recovery_score` as the target variable.
3. Drop identifier/date leakage fields and engineer date month/day-of-year.
4. Preprocess numerical and categorical features with scikit-learn pipelines.
5. Apply imputation, scaling, and one-hot encoding where required.
6. Train baseline, regularized, tree-based, and ensemble regression models.
7. Evaluate with MAE, RMSE, R2, residual plots, actual-vs-predicted plots, cross-validation, and robustness tests.
8. Save tables and visual outputs for RQ1-RQ7.

## Models

Dummy Regressor, Linear Regression, Ridge Regression, ElasticNet, Decision Tree, k-NN, Random Forest, Gradient Boosting, and XGBoost.

## Evaluation Metrics

MAE, RMSE, R2, residual analysis, actual-vs-predicted plots, cross-validation performance, and robustness under noise/missingness.

## Expected Figures And Tables

Baseline performance tables, model comparison tables, preprocessing comparison tables, feature importance tables, model ranking tables, robustness tables, decision matrix, actual-vs-predicted plots, residual plots, feature importance charts, heatmaps, and robustness charts.

## Consistency With Implemented Code

The proposal matches the notebooks, dataset, target variable, models, evaluation metrics, generated result tables, and visual outputs in this repository.
"""
    (ROOT / "Submission_1_Proposal.md").write_text(proposal, encoding="utf-8")
    (ROOT / "requirements.txt").write_text("pandas\nnumpy\nmatplotlib\nseaborn\nscikit-learn\nxgboost\njupyter\nnbformat\nnbclient\npython-docx\n", encoding="utf-8")


def write_docx() -> None:
    ensure_dirs()
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except Exception:
        return

    def cell_shading(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def cell_text(cell, text, bold=False, color=None, size=8):
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(text))
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def fmt(v):
        try:
            x = float(v)
            return f"{x:.3f}" if abs(x) >= 10 else f"{x:.4f}"
        except Exception:
            return str(v)

    def add_df_table(doc, caption, df, cols=None, max_rows=None):
        doc.add_paragraph(caption, style="Caption")
        if cols:
            df = df[cols]
        if max_rows:
            df = df.head(max_rows)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, col in enumerate(df.columns):
            cell_text(table.rows[0].cells[i], col.replace("_", " "), True, (255, 255, 255), 7)
            cell_shading(table.rows[0].cells[i], "1F4E79")
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(df.columns):
                cell_text(cells[i], fmt(row[col]), size=7)
        doc.add_paragraph()

    def add_figure(doc, file, caption):
        file = IMAGE_DIR / file
        if file.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(file), width=Inches(5.8))
            cap = doc.add_paragraph(caption, style="Caption")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    for style in ["Normal", "Title", "Heading 1", "Heading 2", "Caption"]:
        doc.styles[style].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Title"].font.size = Pt(18)
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Caption"].font.size = Pt(9)
    doc.styles["Caption"].font.italic = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Submission 1: Proposal, Code, and GitHub Repository")
    run.bold = True
    run.font.size = Pt(18)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Supervised Learning for WHOOP Fitness Recovery Score Prediction").bold = True

    df = load_data()
    doc.add_heading("Project and Dataset Summary", level=1)
    summary = pd.DataFrame([
        {"Item": "Dataset", "Value": DATASET_NAME},
        {"Item": "Dataset link", "Value": DATASET_LINK},
        {"Item": "Rows", "Value": f"{df.shape[0]:,}"},
        {"Item": "Columns", "Value": df.shape[1]},
        {"Item": "Target variable", "Value": TARGET},
        {"Item": "Task type", "Value": "Regression"},
    ])
    add_df_table(doc, "Table A. Project and dataset summary", summary)
    doc.add_heading("Metric Note", level=1)
    doc.add_paragraph("This is a regression task. The project reports MAE, RMSE, R2, residuals, actual-vs-predicted plots, cross-validation, and robustness outputs.")
    doc.add_heading("Methodology Summary", level=1)
    for item in [
        "Load and inspect the WHOOP Fitness dataset.",
        "Use recovery_score as the target variable.",
        "Preprocess numerical and categorical features with pipelines.",
        "Train baseline, regularized, tree-based, and ensemble regression models.",
        "Save tabular outputs and visual outputs for RQ1-RQ7.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    sections = [
        ("RQ1", "Baseline Performance", "How effectively can baseline supervised learning models predict WHOOP recovery score?", "RQ1_table_report_rounded.csv", ["Model", "MAE", "RMSE", "R2", "Expected_R2_Range", "Matches_Expected_Range"], "RQ1_error_metrics.png", "Figure 1. Baseline model error metrics."),
        ("RQ2", "Model Comparison", "Which supervised learning model achieves the best predictive performance for recovery score?", "RQ2_table_report_rounded.csv", ["Model", "Model_Type", "MAE", "RMSE", "R2", "Rank_by_R2", "Best_Model"], "RQ2_R2_comparison.png", "Figure 2. Candidate model R2 comparison."),
        ("RQ3", "Effect of Preprocessing", "How do preprocessing strategies affect supervised learning performance?", "RQ3_table.csv", None, "RQ3_R2_comparison.png", "Figure 3. Preprocessing strategy comparison."),
        ("RQ4", "Feature Importance and Interpretability", "Which features contribute most to recovery score prediction?", "RQ4_aggregated_feature_importance.csv", ["Rank", "Original_Feature", "Importance", "Importance_Percent"], "RQ4_feature_importance_bar.png", "Figure 4. Top aggregated feature importance."),
        ("RQ5", "Sensitivity to Evaluation Metrics", "How does model ranking change across MAE, RMSE, and R2?", "RQ5_model_metric_rankings.csv", ["Model", "MAE", "RMSE", "R2", "MAE_Rank", "RMSE_Rank", "R2_Rank", "Average_Rank"], "RQ5_ranking_bump_chart.png", "Figure 5. Ranking bump chart."),
        ("RQ6", "Robustness and Generalization", "How robust is the best model under validation, noise, and missingness?", "RQ6_robustness_summary.csv", None, "RQ6_r2_robustness_bar.png", "Figure 6. R2 robustness comparison."),
        ("RQ7", "Practical Usefulness and Final Recommendation", "Which model provides the best practical balance for recovery decision support?", "RQ7_decision_matrix.csv", ["Model", "CV_R2", "Holdout_R2", "Interpretability_Score", "Deployment_Suitability_Score", "Weighted_Decision_Score", "Final_Rank"], "RQ7_decision_matrix_heatmap.png", "Figure 7. Final decision matrix heatmap."),
    ]
    for rq, heading, question, table_file, cols, fig, caption in sections:
        doc.add_heading(f"{rq}. {heading}", level=1)
        doc.add_heading("Research Question", level=2)
        doc.add_paragraph(question)
        doc.add_heading("Actual Results", level=2)
        tdf = pd.read_csv(TABLE_DIR / table_file)
        add_df_table(doc, f"Table {rq}. Actual output table for {heading}", tdf, cols=cols, max_rows=10)
        add_figure(doc, fig, caption)

    doc.add_heading("Final Recommendation", level=1)
    add_df_table(doc, "Table B. Final recommendation", pd.read_csv(TABLE_DIR / "RQ7_final_recommendation.csv"))
    doc.add_heading("Submission Checklist", level=1)
    checklist = pd.DataFrame([
        {"Item": "Proposal file", "Status": "Submission_1_Proposal.docx"},
        {"Item": "README", "Status": "README.md"},
        {"Item": "Code", "Status": "whoop_assignment_pipeline.py and RQ1-RQ7 notebooks"},
        {"Item": "Dataset", "Status": "data/whoop_fitness_dataset_100k.csv"},
        {"Item": "Dataset link", "Status": DATASET_LINK},
        {"Item": "Outputs", "Status": "outputs/tables and outputs/figures"},
    ])
    add_df_table(doc, "Table C. Submission checklist", checklist)
    doc.save(ROOT / "Submission_1_Proposal.docx")


def notebook_code(rq: int) -> str:
    table_map = {
        1: ["RQ1_table_report_rounded.csv", "RQ1_error_metrics.png", "RQ1_actual_vs_predicted.png"],
        2: ["RQ2_table_report_rounded.csv", "RQ2_R2_comparison.png", "RQ2_actual_vs_predicted.png"],
        3: ["RQ3_table.csv", "RQ3_R2_comparison.png", "RQ3_actual_vs_predicted.png"],
        4: ["RQ4_aggregated_feature_importance.csv", "RQ4_feature_importance_bar.png", "RQ4_correlation_heatmap.png"],
        5: ["RQ5_model_metric_rankings.csv", "RQ5_ranking_bump_chart.png", "RQ5_rank_heatmap.png"],
        6: ["RQ6_robustness_summary.csv", "RQ6_r2_robustness_bar.png", "RQ6_rmse_robustness_bar.png"],
        7: ["RQ7_decision_matrix.csv", "RQ7_decision_matrix_heatmap.png", "RQ7_weighted_decision_score.png"],
    }
    table, fig1, fig2 = table_map[rq]
    return f"""from pathlib import Path
import pandas as pd
from IPython.display import Image, display
import sys

ROOT = Path.cwd()
if ROOT.name == 'notebooks':
    ROOT = ROOT.parent
sys.path.insert(0, str(ROOT))

from whoop_assignment_pipeline import DATA_PATH

df = pd.read_csv(DATA_PATH)
print(f"Dataset shape: {{df.shape[0]:,}} rows x {{df.shape[1]}} columns")
display(df.head())

print("This notebook displays the executed RQ{rq} outputs generated by whoop_assignment_pipeline.py.")
print("To regenerate only this RQ, run: from whoop_assignment_pipeline import run_rq{rq}; run_rq{rq}()")

table_path = ROOT / 'outputs' / 'tables' / '{table}'
display(pd.read_csv(table_path).head(15))

for image_name in ['{fig1}', '{fig2}']:
    image_path = ROOT / 'outputs' / 'figures' / 'images' / image_name
    if image_path.exists():
        display(Image(filename=str(image_path)))
"""


def write_notebooks(execute: bool = False) -> None:
    ensure_dirs()
    try:
        import warnings
        import nbformat as nbf
        from nbclient import NotebookClient
    except Exception:
        return

    titles = {
        1: "RQ1: Baseline Performance",
        2: "RQ2: Model Comparison",
        3: "RQ3: Effect of Preprocessing",
        4: "RQ4: Feature Importance and Interpretability",
        5: "RQ5: Sensitivity to Evaluation Metrics",
        6: "RQ6: Robustness and Generalization",
        7: "RQ7: Practical Usefulness and Final Recommendation",
    }
    questions = {
        1: "How effectively can baseline supervised learning models predict WHOOP recovery score?",
        2: "Which supervised learning model achieves the best predictive performance for recovery score?",
        3: "How do preprocessing strategies affect supervised learning performance?",
        4: "Which features contribute most to recovery score prediction?",
        5: "How does model ranking change across MAE, RMSE, and R2?",
        6: "How robust is the best model under validation, noise, and missingness?",
        7: "Which model provides the best practical balance for recovery decision support?",
    }
    for rq in range(1, 8):
        nb = nbf.v4.new_notebook()
        nb["metadata"] = {
            "kernelspec": {
                "display_name": "Python 3",
                "language": "python",
                "name": "python3",
            },
            "language_info": {
                "name": "python",
            },
        }
        nb["cells"] = [
            nbf.v4.new_markdown_cell(f"# {titles[rq]}\n\nResearch question: {questions[rq]}\n\nDataset: WHOOP Fitness Dataset. Target variable: `recovery_score`."),
            nbf.v4.new_code_cell(notebook_code(rq)),
        ]
        path = NOTEBOOK_DIR / f"RQ{rq}_notebook.ipynb"
        if execute:
            with warnings.catch_warnings():
                warnings.filterwarnings(
                    "ignore",
                    message=".*Proactor event loop does not implement add_reader family of methods required for zmq.*",
                    category=RuntimeWarning,
                )
                client = NotebookClient(nb, timeout=1200, kernel_name="python3", resources={"metadata": {"path": str(NOTEBOOK_DIR)}})
                nb = client.execute()
        nbf.write(nb, path)


def run_all() -> None:
    ensure_dirs()
    clean_outputs()
    sns.set_theme(style="whitegrid")
    run_rq1()
    run_rq2()
    run_rq3()
    run_rq4()
    run_rq5()
    run_rq6()
    run_rq7()
    write_markdown_files()
    write_docx()
    write_notebooks(execute=False)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--run-all", action="store_true")
    parser.add_argument("--write-notebooks", action="store_true")
    parser.add_argument("--execute-notebooks", action="store_true")
    args = parser.parse_args()
    if not any(vars(args).values()):
        run_all()
        return
    if args.run_all:
        run_all()
    if args.write_notebooks:
        write_notebooks(execute=args.execute_notebooks)


if __name__ == "__main__":
    main()
